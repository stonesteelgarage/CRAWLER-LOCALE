import os
import re
import json
import subprocess
from pathlib import Path
from datetime import datetime

import pandas as pd
import streamlit as st
from openai import OpenAI
from openpyxl import load_workbook
import pdfplumber
from docx import Document


# ======================================================
# CONFIG IBRIDA: config.py locale + Streamlit secrets
# ======================================================

def get_secret(key, default=None):
    try:
        if key in st.secrets:
            return st.secrets[key]
    except Exception:
        pass

    try:
        import config
        if hasattr(config, key):
            return getattr(config, key)
    except Exception:
        pass

    return default


OPENAI_API_KEY = get_secret("OPENAI_API_KEY")
LOGO_PATH = get_secret("LOGO_PATH", "logo.png")
APP_PASSWORD = get_secret("APP_PASSWORD", "")
MODEL_NAME = get_secret("MODEL_NAME", "gpt-4.1-mini")


# ======================================================
# BASE APP
# ======================================================

APP_NAME = "Vendor Folder Miner"

st.set_page_config(
    page_title=APP_NAME,
    layout="wide",
    initial_sidebar_state="expanded"
)

os.makedirs("output", exist_ok=True)
os.makedirs("logs", exist_ok=True)

if not OPENAI_API_KEY:
    st.error("OPENAI_API_KEY non trovata. Inseriscila in config.py oppure nei secrets di Streamlit.")
    st.stop()

client = OpenAI(api_key=OPENAI_API_KEY)


# ======================================================
# CSS
# ======================================================

st.markdown("""
<style>
.stApp {
    background: linear-gradient(135deg, #071526 0%, #0b2238 45%, #08111f 100%);
    color: white;
}

[data-testid="stSidebar"] {
    background-color: rgba(6, 18, 33, 0.98);
}

h1, h2, h3, h4, h5, h6, p, label, span {
    color: white !important;
}

.stButton > button,
.stDownloadButton > button,
button[kind="primary"],
button[kind="secondary"] {
    background-color: #0b2f4f !important;
    color: white !important;
    border: 1px solid #2aa8d8 !important;
    border-radius: 8px !important;
    font-weight: 700 !important;
}

.stButton > button:hover,
.stDownloadButton > button:hover {
    background-color: #103f68 !important;
    color: white !important;
}

div[data-testid="stMetric"] {
    background-color: rgba(8, 33, 61, 0.85);
    border: 1px solid rgba(0, 210, 255, 0.35);
    border-radius: 14px;
    padding: 18px;
}

.card {
    background: rgba(8, 33, 61, 0.82);
    border: 1px solid rgba(0, 210, 255, 0.25);
    border-radius: 18px;
    padding: 22px;
    margin-bottom: 18px;
}

input[type="password"],
input[type="text"],
textarea {
    background-color: white !important;
    color: black !important;
}

div[data-testid="stAlert"] p {
    color: black !important;
}
</style>
""", unsafe_allow_html=True)


# ======================================================
# LOGIN OPZIONALE
# ======================================================

if APP_PASSWORD:
    if "logged_in" not in st.session_state:
        st.session_state["logged_in"] = False

    if not st.session_state["logged_in"]:
        col1, col2, col3 = st.columns([1, 1.2, 1])
        with col2:
            if os.path.exists(LOGO_PATH):
                st.image(LOGO_PATH, width=160)

            st.markdown(f"""
            <div class="card">
                <h2>{APP_NAME}</h2>
                <p>Accesso riservato</p>
            </div>
            """, unsafe_allow_html=True)

            password = st.text_input("Password", type="password")

            if st.button("Accedi", use_container_width=True):
                if password == APP_PASSWORD:
                    st.session_state["logged_in"] = True
                    st.rerun()
                else:
                    st.error("Password errata")

        st.stop()


# ======================================================
# UTILITY
# ======================================================

SUPPORTED_EXTENSIONS = {
    ".xlsx", ".xlsm", ".csv", ".txt", ".pdf", ".docx"
}


def clean_text(value):
    if value is None:
        return ""

    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass

    value = str(value)
    value = value.replace("\n", " ")
    value = value.replace("\t", " ")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def normalize_supplier_name(name):
    name = clean_text(name).lower()
    name = name.replace(".", "")
    name = name.replace(",", "")
    name = name.replace(" srl", "")
    name = name.replace(" s r l", "")
    name = name.replace(" spa", "")
    name = name.replace(" s p a", "")
    name = name.replace(" società", "")
    name = name.replace(" societa", "")
    name = re.sub(r"\s+", " ", name)
    return name.strip()


def chunk_text(text, max_chars=9000):
    text = clean_text(text)
    chunks = []

    while len(text) > max_chars:
        split_at = text.rfind(" ", 0, max_chars)
        if split_at == -1:
            split_at = max_chars
        chunks.append(text[:split_at])
        text = text[split_at:].strip()

    if text:
        chunks.append(text)

    return chunks


# ======================================================
# PATH WINDOWS / DESKTOP / ONEDRIVE / RETE
# ======================================================

def normalize_user_path(raw_path):
    raw_path = clean_text(raw_path)
    raw_path = raw_path.strip().strip('"').strip("'").strip()

    if not raw_path:
        return Path("")

    if re.match(r"^[A-Za-z]:\\", raw_path):
        if os.name == "nt":
            return Path(raw_path)

        drive = raw_path[0].lower()
        converted = raw_path.replace("\\", "/")
        converted = converted[2:]
        converted = f"/mnt/{drive}{converted}"
        return Path(converted)

    if raw_path.startswith("\\\\"):
        if os.name == "nt":
            return Path(raw_path)
        return Path(raw_path.replace("\\", "/"))

    return Path(raw_path).expanduser()


def possible_fallback_paths(raw_path):
    raw_path = clean_text(raw_path).strip().strip('"').strip("'")
    folder_name = Path(raw_path).name if raw_path else ""

    if not folder_name:
        return []

    home = Path.home()
    candidates = [
        home / "Desktop" / folder_name,
        home / "OneDrive" / "Desktop" / folder_name,
        home / "OneDrive" / folder_name,
        home / "OneDrive - Ghella SpA" / "Desktop" / folder_name,
        home / "OneDrive - Ghella SpA" / folder_name,
    ]

    userprofile = os.environ.get("USERPROFILE", "")
    if userprofile:
        candidates.extend([
            Path(userprofile) / "Desktop" / folder_name,
            Path(userprofile) / "OneDrive" / "Desktop" / folder_name,
            Path(userprofile) / "OneDrive - Ghella SpA" / "Desktop" / folder_name,
            Path(userprofile) / "OneDrive - Ghella SpA" / folder_name,
        ])

    unique = []
    seen = set()
    for c in candidates:
        s = str(c)
        if s not in seen:
            seen.add(s)
            unique.append(c)

    return unique


def resolve_folder_path(raw_path):
    folder = normalize_user_path(raw_path)

    if folder.exists() and folder.is_dir():
        return folder, "OK"

    for candidate in possible_fallback_paths(raw_path):
        if candidate.exists() and candidate.is_dir():
            return candidate, f"Trovata cartella alternativa: {candidate}"

    parent = folder.parent if str(folder) else Path.home()

    if parent.exists():
        try:
            dirs = [x.name for x in parent.iterdir() if x.is_dir()]
            dirs = dirs[:80]
        except Exception:
            dirs = []

        msg = (
            f"Cartella non trovata: {folder}\n\n"
            f"La cartella superiore esiste: {parent}\n\n"
            f"Cartelle presenti lì dentro:\n"
            + "\n".join([f"- {d}" for d in dirs])
        )
        return folder, msg

    msg = (
        f"Cartella non trovata: {folder}\n\n"
        f"Nemmeno la cartella superiore esiste: {parent}\n\n"
        f"Usa il pulsante 'Seleziona cartella dal PC' oppure copia il percorso dalla barra di Esplora File."
    )
    return folder, msg


def get_windows_folder_dialog():
    if os.name != "nt":
        return ""

    ps_script = """
Add-Type -AssemblyName System.Windows.Forms
$dialog = New-Object System.Windows.Forms.FolderBrowserDialog
$dialog.Description = "Seleziona la cartella da analizzare"
$dialog.ShowNewFolderButton = $false
$result = $dialog.ShowDialog()
if ($result -eq [System.Windows.Forms.DialogResult]::OK) {
    Write-Output $dialog.SelectedPath
}
"""

    try:
        completed = subprocess.run(
            ["powershell", "-NoProfile", "-STA", "-Command", ps_script],
            capture_output=True,
            text=True,
            timeout=120
        )

        if completed.returncode == 0:
            return completed.stdout.strip()

    except Exception as e:
        st.warning(f"Selettore cartella Windows non disponibile: {e}")

    return ""


def get_tkinter_folder_dialog():
    try:
        import tkinter as tk
        from tkinter import filedialog

        root = tk.Tk()
        root.withdraw()
        root.attributes("-topmost", True)

        folder = filedialog.askdirectory(title="Seleziona la cartella da analizzare")

        root.destroy()

        return folder or ""

    except Exception:
        return ""


def select_folder_dialog():
    folder = get_windows_folder_dialog()
    if folder:
        return folder

    return get_tkinter_folder_dialog()


def find_files(folder_path):
    folder, diagnostic = resolve_folder_path(folder_path)

    if not folder.exists() or not folder.is_dir():
        raise FileNotFoundError(diagnostic)

    files = []

    for root, dirs, filenames in os.walk(folder):
        for filename in filenames:
            path = Path(root) / filename

            if path.suffix.lower() in SUPPORTED_EXTENSIONS:
                if not filename.startswith("~$"):
                    files.append(path)

    return files, folder, diagnostic


# ======================================================
# LETTURA DOCUMENTI
# ======================================================

def read_excel_file(path):
    texts = []

    try:
        wb = load_workbook(path, read_only=True, data_only=True)
        for ws in wb.worksheets:
            rows_text = []
            max_row = min(ws.max_row or 1, 2500)
            max_col = min(ws.max_column or 1, 80)

            for row in ws.iter_rows(
                min_row=1,
                max_row=max_row,
                max_col=max_col,
                values_only=True
            ):
                cells = [clean_text(v) for v in row if clean_text(v)]
                if cells:
                    rows_text.append(" | ".join(cells))

            if rows_text:
                texts.append(f"FOGLIO: {ws.title}\n" + "\n".join(rows_text))
    except Exception as e:
        texts.append(f"ERRORE LETTURA EXCEL: {e}")

    return "\n\n".join(texts)


def read_csv_file(path):
    try:
        df = pd.read_csv(path, dtype=str, encoding="utf-8", sep=None, engine="python").fillna("")
    except Exception:
        try:
            df = pd.read_csv(path, dtype=str, encoding="latin-1", sep=None, engine="python").fillna("")
        except Exception as e:
            return f"ERRORE LETTURA CSV: {e}"

    return df.to_csv(index=False, sep=" | ")


def read_txt_file(path):
    for enc in ["utf-8", "latin-1", "cp1252"]:
        try:
            return Path(path).read_text(encoding=enc, errors="ignore")
        except Exception:
            continue
    return ""


def read_pdf_file(path):
    texts = []

    try:
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages[:80], start=1):
                txt = page.extract_text() or ""
                if txt.strip():
                    texts.append(f"PAGINA {i}\n{txt}")
    except Exception as e:
        texts.append(f"ERRORE LETTURA PDF: {e}")

    return "\n\n".join(texts)


def read_docx_file(path):
    try:
        doc = Document(path)
        parts = []

        for p in doc.paragraphs:
            txt = clean_text(p.text)
            if txt:
                parts.append(txt)

        for table in doc.tables:
            for row in table.rows:
                cells = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)
    except Exception as e:
        return f"ERRORE LETTURA DOCX: {e}"


def read_any_file(path):
    suffix = path.suffix.lower()

    if suffix in [".xlsx", ".xlsm"]:
        return read_excel_file(path)

    if suffix == ".csv":
        return read_csv_file(path)

    if suffix == ".txt":
        return read_txt_file(path)

    if suffix == ".pdf":
        return read_pdf_file(path)

    if suffix == ".docx":
        return read_docx_file(path)

    return ""


# ======================================================
# OPENAI EXTRACTION
# ======================================================

def extract_suppliers_with_openai(text, source_file):
    chunks = chunk_text(text, max_chars=9000)
    all_suppliers = []

    for idx, chunk in enumerate(chunks[:8], start=1):
        prompt = f"""
Sei un procurement analyst.

Analizza il testo seguente, proveniente da un documento aziendale/procurement.
Devi estrarre tutti i fornitori, imprese, subappaltatori, aziende, produttori, distributori o società citate come soggetti economici.

Per ogni fornitore estrai, quando disponibile:
- supplier_name: nome fornitore / ragione sociale
- vat_number: partita IVA o codice fiscale aziendale
- email
- phone
- address
- website
- offering: cosa offre, cioè servizio/prodotto/materiale/lavorazione collegata
- category: categoria sintetica, es. impianti, carpenteria, cemento, trasporti, consulenza, noleggio, sicurezza, forniture, subappalto, ecc.
- confidence: da 0 a 100

Regole:
- Non inventare dati.
- Se non trovi un campo, lascia stringa vuota.
- Se un nome sembra persona fisica senza azienda, ignoralo.
- Se una società compare senza dettagli ma è chiaramente un fornitore, includila.
- Restituisci SOLO JSON valido.
- Il JSON deve avere questa forma:
{{
  "suppliers": [
    {{
      "supplier_name": "",
      "vat_number": "",
      "email": "",
      "phone": "",
      "address": "",
      "website": "",
      "offering": "",
      "category": "",
      "confidence": 0
    }}
  ]
}}

File origine: {source_file}
Chunk: {idx}/{len(chunks)}

TESTO:
{chunk}
"""

        try:
            response = client.responses.create(
                model=MODEL_NAME,
                input=prompt
            )

            raw = response.output_text.strip()
            raw = raw.replace("```json", "").replace("```", "").strip()
            data = json.loads(raw)

            suppliers = data.get("suppliers", [])
            if isinstance(suppliers, list):
                for s in suppliers:
                    if clean_text(s.get("supplier_name", "")):
                        s["source_file"] = source_file
                        all_suppliers.append(s)

        except Exception as e:
            log_path = Path("logs") / "openai_errors.log"
            with open(log_path, "a", encoding="utf-8") as f:
                f.write(f"\n--- {datetime.now()} | {source_file} | chunk {idx} ---\n{e}\n")

    return all_suppliers


# ======================================================
# DEDUPLICA
# ======================================================

def deduplicate_suppliers(rows):
    merged = {}

    for row in rows:
        name = clean_text(row.get("supplier_name", ""))
        vat = clean_text(row.get("vat_number", ""))
        email = clean_text(row.get("email", ""))
        website = clean_text(row.get("website", ""))

        if vat:
            key = f"vat::{vat.lower()}"
        elif email:
            key = f"email::{email.lower()}"
        elif website:
            key = f"web::{website.lower().replace('https://', '').replace('http://', '').replace('www.', '').strip('/')}"
        else:
            key = f"name::{normalize_supplier_name(name)}"

        if key not in merged:
            merged[key] = {
                "NOME": name,
                "PIVA": vat,
                "EMAIL": email,
                "TELEFONO": clean_text(row.get("phone", "")),
                "INDIRIZZO": clean_text(row.get("address", "")),
                "SITO WEB": website,
                "TIPO SERVIZIO": clean_text(row.get("category", "")),
                "MATERIALE / SERVIZIO OFFERTO": clean_text(row.get("offering", "")),
                "CONFIDENZA AI": row.get("confidence", ""),
                "FILE ORIGINE": clean_text(row.get("source_file", "")),
            }
        else:
            existing = merged[key]

            for target_col, source_key in [
                ("PIVA", "vat_number"),
                ("EMAIL", "email"),
                ("TELEFONO", "phone"),
                ("INDIRIZZO", "address"),
                ("SITO WEB", "website"),
                ("MATERIALE / SERVIZIO OFFERTO", "offering"),
                ("TIPO SERVIZIO", "category"),
            ]:
                current = clean_text(existing.get(target_col, ""))
                new = clean_text(row.get(source_key, ""))

                if new and new not in current:
                    existing[target_col] = f"{current}; {new}".strip("; ")

            source = clean_text(row.get("source_file", ""))
            if source and source not in existing["FILE ORIGINE"]:
                existing["FILE ORIGINE"] += f"; {source}"

            try:
                existing_conf = int(existing.get("CONFIDENZA AI") or 0)
                new_conf = int(row.get("confidence") or 0)
                existing["CONFIDENZA AI"] = max(existing_conf, new_conf)
            except Exception:
                pass

    return list(merged.values())


# ======================================================
# ELABORAZIONE CARTELLA
# ======================================================

def process_folder(folder_path, max_files=None):
    files, resolved_folder, diagnostic = find_files(folder_path)

    if max_files:
        files = files[:max_files]

    all_rows = []
    report_rows = []

    progress = st.progress(0)
    status = st.empty()

    total = len(files)

    if total == 0:
        return files, [], [], None, resolved_folder, diagnostic

    for i, path in enumerate(files, start=1):
        status.info(f"ALMOND Intelligence sta lavorando... File {i}/{total}: {path.name}")

        try:
            text = read_any_file(path)
            chars = len(text)

            if chars < 30:
                report_rows.append({
                    "File": str(path),
                    "Esito": "Saltato",
                    "Dettaglio": "Testo insufficiente"
                })
            else:
                suppliers = extract_suppliers_with_openai(text, str(path))
                all_rows.extend(suppliers)

                report_rows.append({
                    "File": str(path),
                    "Esito": "Analizzato",
                    "Dettaglio": f"Fornitori estratti: {len(suppliers)}"
                })

        except Exception as e:
            report_rows.append({
                "File": str(path),
                "Esito": "Errore",
                "Dettaglio": str(e)
            })

        progress.progress(i / total if total else 1)

    status.success("Analisi cartella terminata")

    deduped = deduplicate_suppliers(all_rows)

    output_path = Path("output") / f"fornitori_estratti_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

    df_suppliers = pd.DataFrame(deduped)
    df_report = pd.DataFrame(report_rows)

    ordered_cols = [
        "NOME",
        "PIVA",
        "EMAIL",
        "TELEFONO",
        "TIPO SERVIZIO",
        "MATERIALE / SERVIZIO OFFERTO",
        "INDIRIZZO",
        "SITO WEB",
        "CONFIDENZA AI",
        "FILE ORIGINE",
    ]

    for col in ordered_cols:
        if col not in df_suppliers.columns:
            df_suppliers[col] = ""

    df_suppliers = df_suppliers[ordered_cols]

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        df_suppliers.to_excel(writer, sheet_name="Fornitori", index=False)
        df_report.to_excel(writer, sheet_name="Log analisi", index=False)

        ws = writer.book["Fornitori"]
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions

        widths = {
            "A": 35,
            "B": 18,
            "C": 35,
            "D": 22,
            "E": 28,
            "F": 55,
            "G": 45,
            "H": 35,
            "I": 16,
            "J": 80,
        }

        for col_letter, width in widths.items():
            ws.column_dimensions[col_letter].width = width

    return files, deduped, report_rows, output_path, resolved_folder, diagnostic


# ======================================================
# SIDEBAR
# ======================================================

with st.sidebar:
    if os.path.exists(LOGO_PATH):
        st.image(LOGO_PATH, width=150)

    st.markdown(f"### {APP_NAME}")
    st.caption("Estrazione fornitori da cartelle locali, Desktop, OneDrive e percorsi di rete")

    st.divider()

    st.markdown("**File supportati**")
    st.write(".xlsx, .xlsm, .csv, .txt, .pdf, .docx")

    st.divider()

    st.markdown("**Configurazione**")
    st.write(f"Modello: `{MODEL_NAME}`")

    st.divider()
    st.caption("Percorsi supportati:")
    st.code(r"C:\Users\nome\Desktop\cartella")
    st.code(r"C:\Users\nome\OneDrive - Ghella SpA\Desktop\cartella")
    st.code(r"\\server\share\cartella")


# ======================================================
# UI PRINCIPALE
# ======================================================

if "folder_path_value" not in st.session_state:
    st.session_state["folder_path_value"] = ""

st.title(APP_NAME)

st.markdown("""
<div class="card">
<h3>Estrazione intelligente fornitori da cartella locale</h3>
<p>
Inserisci o seleziona il percorso di una cartella del PC. L'app scandaglia sottocartelle e documenti,
legge Excel, PDF, Word, CSV e TXT, usa OpenAI per individuare fornitori, contatti e servizi,
e produce un unico file Excel finale.
</p>
</div>
""", unsafe_allow_html=True)

col_select_1, col_select_2 = st.columns([1, 1])

with col_select_1:
    if st.button("📁 Seleziona cartella dal PC", use_container_width=True):
        selected_folder = select_folder_dialog()
        if selected_folder:
            st.session_state["folder_path_value"] = selected_folder
            st.rerun()
        else:
            st.warning("Nessuna cartella selezionata. Se il popup non si apre, incolla manualmente il percorso.")

with col_select_2:
    if st.button("🧹 Svuota percorso", use_container_width=True):
        st.session_state["folder_path_value"] = ""
        st.rerun()

folder_path = st.text_input(
    "Percorso cartella locale o percorso di rete",
    value=st.session_state["folder_path_value"],
    placeholder=r"C:\Users\samandorla\OneDrive - Ghella SpA\Desktop\DB COLMETO oppure \\server\share\cartella"
)

st.session_state["folder_path_value"] = folder_path

with st.expander("Verifica percorso prima dell'analisi"):
    if folder_path:
        resolved, diagnostic = resolve_folder_path(folder_path)

        st.write("Percorso inserito:")
        st.code(folder_path)

        st.write("Percorso interpretato da Python:")
        st.code(str(resolved))

        st.write(f"Cartella esistente: {resolved.exists()}")
        st.write(f"È una cartella: {resolved.is_dir()}")
        st.write(f"Cartella superiore: {resolved.parent}")
        st.write(f"Cartella superiore esistente: {resolved.parent.exists()}")

        if resolved.exists() and resolved.is_dir():
            st.success(f"Cartella trovata: {resolved}")
            try:
                supported_count = 0
                for root, dirs, filenames in os.walk(resolved):
                    for filename in filenames:
                        if Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS and not filename.startswith('~$'):
                            supported_count += 1
                st.info(f"File supportati trovati nella cartella e sottocartelle: {supported_count}")
            except Exception as e:
                st.warning(f"Cartella trovata, ma non riesco a contare i file: {e}")
        else:
            st.error(diagnostic)
    else:
        st.info("Inserisci o seleziona una cartella.")

col_a, col_b = st.columns(2)

with col_a:
    limit_files = st.checkbox("Limita numero file per test", value=True)

with col_b:
    max_files = st.number_input("Numero massimo file", min_value=1, max_value=5000, value=100, step=10)

effective_max_files = int(max_files) if limit_files else None

if st.button("Avvia analisi cartella", use_container_width=True):
    if not folder_path:
        st.warning("Inserisci o seleziona il percorso della cartella.")
    else:
        try:
            files, suppliers, report_rows, output_path, resolved_folder, diagnostic = process_folder(folder_path, effective_max_files)

            st.success("Analisi completata")
            st.info(f"Cartella usata: {resolved_folder}")
            st.info(f"File analizzati: {len(files)}")
            st.info(f"Fornitori univoci trovati: {len(suppliers)}")

            if diagnostic != "OK":
                st.caption(diagnostic)

            if suppliers and output_path:
                st.subheader("Anteprima fornitori estratti")
                st.dataframe(pd.DataFrame(suppliers), use_container_width=True)

                with open(output_path, "rb") as f:
                    st.download_button(
                        "Scarica Excel fornitori",
                        data=f,
                        file_name=output_path.name,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
            else:
                st.warning("Nessun fornitore trovato nei documenti analizzati.")

            with st.expander("Log analisi file"):
                st.dataframe(pd.DataFrame(report_rows), use_container_width=True)

        except Exception as e:
            st.error(str(e))
